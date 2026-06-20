import io
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def generar_jd_word_doc(busqueda):
    """
    Genera el documento de Job Description en formato Word .docx.
    Retorna un buffer io.BytesIO con el archivo guardado.
    """
    doc = Document()
    
    # Configuración de márgenes premium
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Definición de estilo general (12pt Arial para máxima fidelidad a los ejemplos)
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)
    font.color.rgb = RGBColor(0x21, 0x25, 0x29) # Charcoal oscuro premium
    
    # Encabezado Corporativo sutil
    p_header = doc.add_paragraph()
    p_header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_header.paragraph_format.space_after = Pt(24)
    run_header = p_header.add_run("ConfiaRH - Consultora de Recursos Humanos")
    run_header.font.size = Pt(8.5)
    run_header.font.italic = True
    run_header.font.color.rgb = RGBColor(0x77, 0x77, 0x77)
    
    # Título Principal
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_after = Pt(18)
    run_title = p_title.add_run(busqueda.titulo_puesto)
    run_title.bold = True
    run_title.font.size = Pt(16)
    run_title.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A) # Slate 900
    
    # Función auxiliar para agregar líneas de Ficha Técnica
    def agregar_ficha_linea(label, value, is_bold_val=False, color_val=None):
        if not value:
            return
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        run_lbl = p.add_run(f"{label}: ")
        run_lbl.bold = True
        run_val = p.add_run(str(value))
        run_val.bold = is_bold_val
        if color_val:
            run_val.font.color.rgb = color_val

    # Renderizar detalles de ficha técnica
    agregar_ficha_linea("Empresa", busqueda.empresa.nombre)
    
    fecha_ini_str = "A convenir"
    if busqueda.fecha_inicio:
        fecha_ini_str = busqueda.fecha_inicio.strftime("%d/%m/%Y")
    agregar_ficha_linea("Fecha de inicio", fecha_ini_str)
    
    agregar_ficha_linea("Vacantes Disponibles", busqueda.cantidad_posiciones)
    
    # Salario formateado
    salario_str = ""
    if busqueda.salario_minimo > 0 and busqueda.salario_maximo > 0:
        salario_str = f"${busqueda.salario_minimo:,.0f} a ${busqueda.salario_maximo:,.0f}"
    elif busqueda.salario_minimo > 0:
        salario_str = f"${busqueda.salario_minimo:,.0f} Netos"
    else:
        salario_str = "A convenir / No especificado"
    agregar_ficha_linea("Rango Salarial", salario_str)
    
    agregar_ficha_linea("Zona", busqueda.ubicacion_puesto or "No especificada")
    
    # Observaciones destacadas
    if busqueda.observaciones:
        agregar_ficha_linea("Observaciones", busqueda.observaciones, is_bold_val=True, color_val=RGBColor(0xC0, 0x39, 0x2B))
        
    def agregar_titulo_bloque(titulo):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(titulo)
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

    # Misión o Descripción del puesto
    if busqueda.mision_puesto or busqueda.descripcion:
        agregar_titulo_bloque("Descripción del puesto:")
        p_desc = doc.add_paragraph(busqueda.mision_puesto or busqueda.descripcion)
        p_desc.paragraph_format.space_after = Pt(8)

    # Responsabilidades
    if busqueda.responsabilidades:
        agregar_titulo_bloque("Responsabilidades:")
        for resp in busqueda.responsabilidades.split('\n'):
            resp = resp.strip()
            if resp:
                resp_cleaned = resp.lstrip('-•* ').strip()
                p_item = doc.add_paragraph(resp_cleaned, style='List Bullet')
                p_item.paragraph_format.space_after = Pt(3)

    # Requisitos
    if busqueda.requisitos_excluyentes:
        agregar_titulo_bloque("Requisitos:")
        for req in busqueda.requisitos_excluyentes.split('\n'):
            req = req.strip()
            if req:
                req_cleaned = req.lstrip('-•* ').strip()
                p_item = doc.add_paragraph(req_cleaned, style='List Bullet')
                p_item.paragraph_format.space_after = Pt(3)

    # Preguntas para incluir en el informe
    if busqueda.preguntas_informe:
        agregar_titulo_bloque("PREGUNTAS PARA INCLUIR EN EL INFORME (Deberán ser contestadas correctamente para presentar un candidato)")
        for preg in busqueda.preguntas_informe.split('\n'):
            preg = preg.strip()
            if preg:
                preg_cleaned = preg.lstrip('-•* ').strip()
                p_item = doc.add_paragraph(preg_cleaned)
                p_item.paragraph_format.space_after = Pt(4)

    # Candidato Ideal
    if busqueda.candidato_ideal:
        agregar_titulo_bloque("CANDIDATO IDEAL:")
        p_ideal = doc.add_paragraph(busqueda.candidato_ideal)
        p_ideal.paragraph_format.space_after = Pt(8)

    # Beneficios
    if busqueda.beneficios:
        agregar_titulo_bloque("Beneficios:")
        for ben in busqueda.beneficios.split('\n'):
            ben = ben.strip()
            if ben:
                ben_cleaned = ben.lstrip('-•* ').strip()
                p_item = doc.add_paragraph(ben_cleaned, style='List Bullet')
                p_item.paragraph_format.space_after = Pt(3)

    # Bloque avanzado: Descripción Puesto / Ficha Operativa
    agregar_titulo_bloque("  Descripción Puesto")
    
    agregar_ficha_linea("Ubicación", busqueda.ubicacion_puesto)
    agregar_ficha_linea("Horario", busqueda.jornada_laboral)
    
    if busqueda.descansos:
        agregar_ficha_linea("Descansos", busqueda.descansos)
        
    agregar_ficha_linea("Beneficios", salario_str)
    
    if busqueda.convenio:
        agregar_ficha_linea("Convenio", busqueda.convenio)
        
    agregar_ficha_linea("Experiencia requerida", f"Mínimo {busqueda.experiencia_minima_anios} años en posiciones similares" if busqueda.experiencia_minima_anios else None)
    
    if busqueda.edad_rango:
        agregar_ficha_linea("Edad", busqueda.edad_rango)
        
    if busqueda.zona_residencia:
        agregar_ficha_linea("Zona de residencia", busqueda.zona_residencia)
        
    if busqueda.requisitos_deseables:
        deseables_limpio = busqueda.requisitos_deseables.replace("\n", ", ").lstrip('-•* ')
        agregar_ficha_linea("Valorado", deseables_limpio)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
